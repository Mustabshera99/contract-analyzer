"""
Document processing service for contract analysis.

This module provides utilities for processing various document formats,
extracting text content, and preparing documents for analysis.
"""

import logging
import os
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union

import structlog

try:
	import magic

	MAGIC_AVAILABLE = True
except ImportError:
	MAGIC_AVAILABLE = False
	magic = None

from docx import Document

try:
	from unstructured.partition.docx import partition_docx
	from unstructured.partition.pdf import partition_pdf

	UNSTRUCTURED_AVAILABLE = True
except ImportError:
	UNSTRUCTURED_AVAILABLE = False
	partition_docx = None
	partition_pdf = None

from ..core.exceptions import DocumentProcessingError, ValidationError

logger = logging.getLogger(__name__)
structured_logger = structlog.get_logger(__name__)


@dataclass
class ProcessedDocument:
	"""Container for processed document data."""

	content: str
	filename: str
	file_type: str
	file_size: int
	metadata: Dict[str, Union[str, int, float]]


class DocumentValidator:
	"""Validates uploaded documents before processing."""

	SUPPORTED_FORMATS = {".pdf", ".docx", ".txt"}
	MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

	@classmethod
	def validate_file(cls, file_path: Path, original_filename: str) -> Tuple[bool, Optional[str]]:
		"""
		Validate uploaded file format and size.

		Args:
		    file_path: Path to the uploaded file
		    original_filename: Original filename from upload

		Returns:
		    Tuple of (is_valid, error_message)
		"""
		try:
			# Check file exists
			if not file_path.exists():
				return False, "File not found"

			# Check file size
			file_size = file_path.stat().st_size
			if file_size > cls.MAX_FILE_SIZE:
				return False, f"File size ({file_size} bytes) exceeds maximum allowed size ({cls.MAX_FILE_SIZE} bytes)"

			if file_size == 0:
				return False, "File is empty"

			# Check file extension
			file_extension = Path(original_filename).suffix.lower()
			if file_extension not in cls.SUPPORTED_FORMATS:
				return False, f"Unsupported file format: {file_extension}. Supported formats: {', '.join(cls.SUPPORTED_FORMATS)}"

			# Validate MIME type using python-magic if available
			if MAGIC_AVAILABLE:
				try:
					mime_type = magic.from_file(str(file_path), mime=True)
					expected_mime_types = {
						".pdf": ["application/pdf"],
						".docx": ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
						".txt": ["text/plain"],
					}

					if mime_type not in expected_mime_types.get(file_extension, []):
						return False, f"File content does not match extension. Expected {file_extension}, got MIME type: {mime_type}"

				except Exception as e:
					logger.warning(f"Could not validate MIME type: {e}")
					# Continue without MIME validation if magic fails
			else:
				logger.info("python-magic not available, skipping MIME type validation")

			return True, None

		except Exception as e:
			logger.error(f"File validation error: {e}")
			return False, f"File validation failed: {e!s}"


class DocumentProcessor:
	"""Main document processing class."""

	def __init__(self):
		self.validator = DocumentValidator()

	def process_document(self, file_path: Path, original_filename: str) -> ProcessedDocument:
		"""
		Process a document file and extract text content.

		Args:
		    file_path: Path to the document file
		    original_filename: Original filename from upload

		Returns:
		    ProcessedDocument with extracted content and metadata

		Raises:
		    ValidationError: If file validation fails
		    DocumentProcessingError: If document processing fails
		"""
		# Validate file
		is_valid, error_message = self.validator.validate_file(file_path, original_filename)
		if not is_valid:
			raise ValidationError(f"File validation failed: {error_message}")

		file_extension = Path(original_filename).suffix.lower()
		file_size = file_path.stat().st_size

		try:
			# Extract text based on file type
			if file_extension == ".pdf":
				content, metadata = self._process_pdf(file_path)
			elif file_extension == ".docx":
				content, metadata = self._process_docx(file_path)
			elif file_extension == ".txt":
				content, metadata = self._process_txt(file_path)
			else:
				raise DocumentProcessingError(f"Unsupported file type: {file_extension}")

			# Clean and normalize content
			cleaned_content = self._clean_and_normalize_text(content)

			# Add processing metadata
			metadata.update(
				{"original_length": len(content), "cleaned_length": len(cleaned_content), "processing_method": f"{file_extension[1:]}_extraction"}
			)

			return ProcessedDocument(
				content=cleaned_content, filename=original_filename, file_type=file_extension, file_size=file_size, metadata=metadata
			)

		except Exception as e:
			logger.error(f"Document processing failed for {original_filename}: {e}")
			raise DocumentProcessingError(f"Failed to process document: {e!s}")

	def _process_pdf(self, file_path: Path) -> Tuple[str, Dict]:
		"""Extract text from PDF file using pypdf or unstructured library."""
		# Try pypdf first (lighter dependency)
		try:
			import pypdf
			return self._process_pdf_pypdf(file_path)
		except ImportError:
			pass

		# Fallback to unstructured if available
		if not UNSTRUCTURED_AVAILABLE:
			raise DocumentProcessingError("PDF processing requires pypdf or unstructured library. Please install with: pip install pypdf")

		try:
			# Use unstructured to partition PDF
			elements = partition_pdf(str(file_path))

			# Extract text from elements
			text_content = []
			metadata = {"total_elements": len(elements), "element_types": {}}

			for element in elements:
				if hasattr(element, "text") and element.text.strip():
					text_content.append(element.text.strip())

				# Track element types for metadata
				element_type = type(element).__name__
				metadata["element_types"][element_type] = metadata["element_types"].get(element_type, 0) + 1

			full_text = "\n\n".join(text_content)

			if not full_text.strip():
				raise DocumentProcessingError("No text content extracted from PDF")

			return full_text, metadata

		except Exception as e:
			logger.error(f"PDF processing error: {e}")
			raise DocumentProcessingError(f"Failed to extract text from PDF: {e!s}")

	def _process_pdf_pypdf(self, file_path: Path) -> Tuple[str, Dict]:
		"""Extract text from PDF using pypdf (lighter alternative)."""
		try:
			import pypdf
			
			text_content = []
			metadata = {"pages_count": 0, "extraction_method": "pypdf"}
			
			with open(file_path, 'rb') as file:
				pdf_reader = pypdf.PdfReader(file)
				metadata["pages_count"] = len(pdf_reader.pages)
				
				for page_num, page in enumerate(pdf_reader.pages):
					try:
						page_text = page.extract_text()
						if page_text.strip():
							text_content.append(page_text.strip())
					except Exception as e:
						logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
						continue
			
			full_text = "\n\n".join(text_content)
			
			if not full_text.strip():
				raise DocumentProcessingError("No text content extracted from PDF")
			
			return full_text, metadata
			
		except Exception as e:
			logger.error(f"pypdf processing error: {e}")
			raise DocumentProcessingError(f"Failed to extract text from PDF with pypdf: {e!s}")

	def _process_docx(self, file_path: Path) -> Tuple[str, Dict]:
		"""Extract text from DOCX file with formatting preservation."""
		try:
			# Use unstructured for primary extraction if available
			if UNSTRUCTURED_AVAILABLE:
				elements = partition_docx(str(file_path))

				# Extract text from elements
				text_content = []
				metadata = {"total_elements": len(elements), "element_types": {}}

				for element in elements:
					if hasattr(element, "text") and element.text.strip():
						text_content.append(element.text.strip())

					# Track element types for metadata
					element_type = type(element).__name__
					metadata["element_types"][element_type] = metadata["element_types"].get(element_type, 0) + 1

				full_text = "\n\n".join(text_content)

				# Fallback to python-docx if unstructured fails or returns empty
				if not full_text.strip():
					logger.warning("Unstructured extraction failed, falling back to python-docx")
					full_text, docx_metadata = self._process_docx_fallback(file_path)
					metadata.update(docx_metadata)
			else:
				# Use python-docx directly if unstructured is not available
				logger.info("Unstructured not available, using python-docx for DOCX processing")
				full_text, metadata = self._process_docx_fallback(file_path)

			if not full_text.strip():
				raise DocumentProcessingError("No text content extracted from DOCX")

			return full_text, metadata

		except Exception as e:
			logger.error(f"DOCX processing error: {e}")
			raise DocumentProcessingError(f"Failed to extract text from DOCX: {e!s}")

	def _process_docx_fallback(self, file_path: Path) -> Tuple[str, Dict]:
		"""Fallback DOCX processing using python-docx."""
		try:
			doc = Document(str(file_path))

			paragraphs = []
			for paragraph in doc.paragraphs:
				if paragraph.text.strip():
					paragraphs.append(paragraph.text.strip())

			# Extract text from tables
			table_content = []
			for table in doc.tables:
				for row in table.rows:
					row_text = []
					for cell in row.cells:
						if cell.text.strip():
							row_text.append(cell.text.strip())
					if row_text:
						table_content.append(" | ".join(row_text))

			# Combine paragraphs and tables
			all_content = paragraphs + table_content
			full_text = "\n\n".join(all_content)

			metadata = {"paragraphs_count": len(paragraphs), "tables_count": len(doc.tables), "extraction_method": "python-docx_fallback"}

			return full_text, metadata

		except Exception as e:
			logger.error(f"DOCX fallback processing error: {e}")
			raise DocumentProcessingError(f"Fallback DOCX processing failed: {e!s}")

	def _process_txt(self, file_path: Path) -> Tuple[str, Dict]:
		"""Extract text from TXT file."""
		try:
			# Read text file with encoding detection
			import chardet

			with open(file_path, "rb") as f:
				raw_data = f.read()

			# Detect encoding
			encoding_result = chardet.detect(raw_data)
			encoding = encoding_result.get("encoding", "utf-8")
			confidence = encoding_result.get("confidence", 0)

			# Try to decode with detected encoding
			try:
				content = raw_data.decode(encoding)
			except (UnicodeDecodeError, LookupError):
				# Fallback to utf-8 with error handling
				content = raw_data.decode("utf-8", errors="replace")
				encoding = "utf-8"
				confidence = 0.5

			# Create metadata
			metadata = {
				"encoding": encoding,
				"encoding_confidence": confidence,
				"extraction_method": "text_file_direct",
				"lines_count": len(content.splitlines()),
			}

			if not content.strip():
				raise DocumentProcessingError("Text file is empty or contains only whitespace")

			return content, metadata

		except Exception as e:
			logger.error(f"TXT processing error: {e}")
			raise DocumentProcessingError(f"Failed to extract text from TXT file: {e!s}")

	def _clean_and_normalize_text(self, text: str) -> str:
		"""Clean and normalize extracted text content."""
		if not text:
			return ""

		# Remove excessive whitespace
		text = re.sub(r"\s+", " ", text)

		# Remove excessive newlines (more than 2 consecutive)
		text = re.sub(r"\n{3,}", "\n\n", text)

		# Clean up common OCR artifacts
		text = re.sub(r"[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}\"\'\/\\\@\#\$\%\^\&\*\+\=\<\>\|]", " ", text)

		# Normalize quotes
		text = re.sub(r'["""]', '"', text)
		text = re.sub(r"[''']", "'", text)

		# Remove leading/trailing whitespace from each line
		lines = [line.strip() for line in text.split("\n")]
		text = "\n".join(line for line in lines if line)

		# Final cleanup
		text = text.strip()

		return text


class DocumentProcessingService:
	"""High-level service for document processing operations."""

	def __init__(self):
		self.processor = DocumentProcessor()

	def process_document(self, file_path: str, filename: str) -> ProcessedDocument:
		"""
		Process a document synchronously.

		Args:
			file_path: Path to the file to process
			filename: Original filename

		Returns:
			ProcessedDocument with extracted content
		"""
		from pathlib import Path

		return self.processor.process_document(Path(file_path), filename)

	async def process_uploaded_file(self, file_content: bytes, filename: str) -> ProcessedDocument:
		"""
		Process an uploaded file from bytes content.

		Args:
		    file_content: Raw file bytes
		    filename: Original filename

		Returns:
		    ProcessedDocument with extracted content
		"""
		file_type = Path(filename).suffix.lower()
		file_size = len(file_content)

		# Use structured logging
		structured_logger.info("Starting document processing", filename=filename, file_type=file_type, file_size=file_size)

		# Import monitoring components
		from ..core.monitoring import performance_monitor

		# Create temporary file
		with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as temp_file:
			temp_file.write(file_content)
			temp_file_path = Path(temp_file.name)

		try:
			# Process the document with monitoring
			async with performance_monitor.monitor_document_processing(file_type, filename):
				result = self.processor.process_document(temp_file_path, filename)

				# Log successful processing
				structured_logger.info(
					"Document processing completed",
					filename=filename,
					file_type=file_type,
					content_length=len(result.content),
					processing_method=result.metadata.get("processing_method"),
				)

				return result

		except Exception as e:
			# Log processing error
			structured_logger.error("Document processing failed", filename=filename, file_type=file_type, error=str(e))
			raise

		finally:
			# Clean up temporary file
			try:
				temp_file_path.unlink()
			except Exception as e:
				logger.warning(f"Failed to clean up temporary file {temp_file_path}: {e}")

	def get_supported_formats(self) -> List[str]:
		"""Get list of supported file formats."""
		return list(DocumentValidator.SUPPORTED_FORMATS)

	def get_max_file_size(self) -> int:
		"""Get maximum allowed file size in bytes."""
		return DocumentValidator.MAX_FILE_SIZE
